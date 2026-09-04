"""Document loaders for various file formats."""

import os
from pathlib import Path

from src.core.constants import DocumentType


def load_single_file(file_path: Path) -> str:
    """Load content from a single file."""
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = file_path.suffix.lower()
    if suffix in (".md", ".txt"):
        return file_path.read_text(encoding="utf-8")
    elif suffix == ".pdf":
        return _load_pdf(file_path)
    elif suffix == ".docx":
        return _load_docx(file_path)
    elif suffix == ".xlsx":
        return _load_xlsx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def _load_pdf(file_path: Path) -> str:
    """Extract text from a PDF file using pypdf."""
    from pypdf import PdfReader
    reader = PdfReader(str(file_path))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def _load_docx(file_path: Path) -> str:
    """Extract text from a DOCX file using python-docx."""
    from docx import Document
    doc = Document(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n\n".join(paragraphs)


def _load_xlsx(file_path: Path) -> str:
    """Extract text from an XLSX file using openpyxl."""
    from openpyxl import load_workbook
    wb = load_workbook(str(file_path), read_only=True, data_only=True)
    sheets_text = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            header = f"Sheet: {sheet_name}"
            body = "\n".join(rows)
            sheets_text.append(header + "\n" + body)
    wb.close()
    return "\n\n".join(sheets_text)


def _count_lines_before(text: str, chunk_start: int) -> int:
    """Count the line number where a character offset occurs."""
    return text[:chunk_start].count("\n") + 1


def load_documents_with_locations(data_dir: Path) -> list[dict]:
    """Load documents with line/page location tracking for citations."""
    documents = []
    supported_suffixes = {".md", ".txt", ".pdf", ".docx", ".xlsx"}

    for root, _dirs, files in os.walk(data_dir):
        for file in sorted(files):
            file_path = Path(root) / file
            if file_path.suffix.lower() not in supported_suffixes:
                continue

            doc_type = _infer_doc_type(file_path)
            suffix = file_path.suffix.lower()

            if suffix == ".pdf":
                locations = _load_pdf_with_pages(file_path)
            elif suffix in (".docx", ".xlsx"):
                # Non-PDF formats: single "page" with full content
                content = load_single_file(file_path)
                locations = [{"text": content, "page": 1, "line": 1}]
            else:
                locations = _load_text_with_lines(file_path)

            content = "\n\n".join(loc["text"] for loc in locations if loc["text"].strip())

            documents.append({
                "content": content,
                "metadata": {
                    "source": str(file_path),
                    "filename": file_path.name,
                },
                "locations": locations,
                "doc_type": doc_type.value,
            })

    return documents


def _load_pdf_with_pages(file_path: Path) -> list[dict]:
    """Load PDF with page number and paragraph-level line tracking."""
    from pypdf import PdfReader
    reader = PdfReader(str(file_path))
    locations = []
    for page_num, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        if not text.strip():
            continue
        paragraphs = text.split("\n\n")
        non_empty = [p for p in paragraphs if p.strip()]
        if len(non_empty) <= 1 and "\n" in text:
            paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
        para_line = 1
        for para in paragraphs:
            if not para.strip():
                continue
            locations.append({
                "text": para.strip(),
                "page": page_num,
                "line": para_line,
            })
            para_line += para.count("\n") + 1
    return locations


def _load_text_with_lines(file_path: Path) -> list[dict]:
    """Load text file with line number tracking."""
    content = file_path.read_text(encoding="utf-8")
    paragraphs = content.split("\n\n")
    locations = []
    line_num = 1
    for para in paragraphs:
        locations.append({
            "text": para,
            "page": 1,
            "line": line_num,
        })
        line_num += para.count("\n") + 2
    return locations


def _infer_doc_type(file_path: Path) -> DocumentType:
    """Infer document type from file path or content."""
    path_str = str(file_path).lower()
    if "term_sheet" in path_str:
        return DocumentType.TERM_SHEET
    elif "financial_model" in path_str:
        return DocumentType.FINANCIAL_MODEL
    elif "investment_memo" in path_str:
        return DocumentType.INVESTMENT_MEMO
    elif "compliance" in path_str:
        return DocumentType.COMPLIANCE_DOC
    elif "portfolio" in path_str:
        return DocumentType.PORTFOLIO_REPORT
    else:
        return DocumentType.LEGAL_AGREEMENT


def load_documents(data_dir: Path) -> list[dict]:
    """Load all documents from a directory tree with location tracking."""
    return load_documents_with_locations(data_dir)
