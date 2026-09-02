"""Tests for multi-format document loading."""

import tempfile
from pathlib import Path

from src.ingestion.loader import load_single_file


def test_load_txt():
    with tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False) as f:
        f.write("Hello world\nSecond line")
        f.flush()
        content = load_single_file(Path(f.name))
        assert "Hello world" in content


def test_load_md():
    with tempfile.NamedTemporaryFile(suffix=".md", mode="w", delete=False) as f:
        f.write("# Title\n\nBody text")
        f.flush()
        content = load_single_file(Path(f.name))
        assert "Title" in content


def test_load_docx():
    from docx import Document
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc = Document()
        doc.add_heading("Test Document", level=1)
        doc.add_paragraph("This is test content.")
        doc.save(f.name)
        content = load_single_file(Path(f.name))
        assert "Test Document" in content
        assert "test content" in content


def test_load_xlsx():
    from openpyxl import Workbook
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
        wb = Workbook()
        ws = wb.active
        ws["A1"] = "Company"
        ws["B1"] = "Revenue"
        ws["A2"] = "Acme"
        ws["B2"] = 1000000
        wb.save(f.name)
        content = load_single_file(Path(f.name))
        assert "Acme" in content
        assert "1000000" in content


def test_unsupported_format():
    import pytest
    with tempfile.NamedTemporaryFile(suffix=".xyz", mode="w", delete=False) as f:
        f.write("data")
        f.flush()
        with pytest.raises(ValueError, match="Unsupported"):
            load_single_file(Path(f.name))
